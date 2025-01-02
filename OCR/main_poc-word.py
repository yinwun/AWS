import os
from docx import Document
from PIL import Image
import boto3

def extract_text_and_images_from_docx(docx_path, output_folder):
    """
    提取 .docx 文件中的文字和图片，并按顺序返回。
    图片会保存到指定文件夹。
    """
    doc = Document(docx_path)
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    content = []  # 保存按顺序提取的内容

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # 提取段落文字
            content.append({"type": "text", "content": paragraph.text.strip()})

    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.target_ref:  # 提取图片
            img_data = rel.target_part.blob
            img_path = os.path.join(output_folder, f"image_{i + 1}.png")
            with open(img_path, "wb") as img_file:
                img_file.write(img_data)
            content.append({"type": "image", "path": img_path})

    return content

def extract_text_with_textract(image_path):
    """
    使用 AWS Textract 对图片进行 OCR。
    """
    textract = boto3.client('textract', region_name='us-east-1')
    try:
        with open(image_path, 'rb') as img_file:
            img_bytes = img_file.read()

        response = textract.detect_document_text(Document={'Bytes': img_bytes})

        # 提取文字内容
        text_blocks = response.get('Blocks', [])
        extracted_text = [block['Text'] for block in text_blocks if block['BlockType'] == 'LINE']
        return '\n'.join(extracted_text)
    except Exception as e:
        print(f"Error processing {image_path}: {e}")
        return ""

def process_docx_with_textract(docx_path):
    """
    主函数：提取 .docx 文件中的文字和图片，并进行 OCR。
    """
    output_folder = "extracted_images"
    content = extract_text_and_images_from_docx(docx_path, output_folder)

    if not content:
        print("No content found in the document.")
        return

    answers = []
    results = []

    print("Processing document content...\n")
    for item in content:
        if item["type"] == "text":
            print("Text Content:")
            print(item["content"])
            print("-" * 50)
            answers.append(item["content"])
        elif item["type"] == "image":
            print(f"Image found: {item['path']}")
            print("Performing OCR...")
            text = extract_text_with_textract(item["path"])
            print("OCR Result:")
            print(text)
            print("-" * 50)
            results.append(text)
    
    write_to_txt(answers, results)
    
def write_to_word(answers, results):
    # 创建一个新的 Word 文档
    doc = Document()

    # 遍历 A 和 B 数组，组合问题和答案并添加到文档
    for i in range(len(answers)):
        question = results[i].split('\n', 1)[-1]
        answer = answers[i]
        
        # 添加问题的序号和内容到 Word 文档
        doc.add_paragraph(f"{i+1}. {question}")
        
        # 添加答案到 Word 文档
        doc.add_paragraph(f"答案: {answer}")
        
        # 添加空行分隔
        doc.add_paragraph("\n")

    # 保存 Word 文档
    output_word_path = 'questions_and_answers_with_numbers.docx'
    doc.save(output_word_path)

    print(f"Word document saved at {output_word_path}")

def write_to_txt(answers, results):
    # 创建或打开一个文本文件以写入
    output_txt_path = 'questions_and_answers_with_numbers.txt'
    x = 0
    try:
        with open(output_txt_path, 'w', encoding='utf-8') as file:
            # 遍历 answers 和 results 数组，组合问题和答案并写入文件
            for i in range(len(answers)):
                number = results[i].split('\n', 1)[0].replace('Question #', '').strip()
                #print(number)
                question = results[i].split('\n', 1)[-1]  # 去掉第一个换行符前的内容
                answer = ''
                try:
                    number_int = int(number)
                    answer = answers[number_int - 1]
                except:
                    print(f"Error: answers[{number}] is out of range.")
                    
                
                # 写入问题的序号和内容到文本文件
                file.write(f"{number}. {question}\n")
                
                # 写入答案到文本文件
                file.write(f"答案: {answer}\n")
                
                # 添加空行分隔
                file.write("\n")

                x = x + 1
                print(x)
        
        print(f"Text file successfully saved at {output_txt_path}")
    
    except Exception as e:
        print(f"An error occurred while writing to the file: {e}")

# 示例文档路径
word_path = r'E:\Sandbox\github\poc\img\CAD_Answer.docx'

# 运行
process_docx_with_textract(word_path)
